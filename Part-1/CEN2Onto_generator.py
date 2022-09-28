from docx import Document
from rdflib.namespace import FOAF
from rdflib import URIRef, Graph, Literal, RDF, RDFS, Namespace
import re


def create_properties(table, sup_prop_name):

    # create super property of all the properties in this table
    sup_prop = URIRef(default_ns + sup_prop_name)
    store.add((sup_prop, RDF.type, RDF.Property))

    # go through each row in the table
    # skip the header row
    for row in table.rows[1:]:
        # Create an identifier to the property which is the GUID
        prop = URIRef(default_ns + Literal(row.cells[0].text.strip()))

        # Add triples using store's add method.
        store.add((prop, RDF.type, RDF.Property))

        store.add((prop, RDFS.subPropertyOf, sup_prop))

        store.add((prop, prop_id, Literal(row.cells[1].text.strip())))
        store.add((prop, RDFS.label, Literal(row.cells[2].text.strip())))
        store.add((prop, RDFS.comment, Literal(row.cells[3].text.strip())))


# read the ISO standard document
file_name = "C:/Users/kebreh/Fagerhult/Generate_Ontology/BIM_Properties.docx"
document = Document(file_name)

# creat a triple store
store = Graph()
# bind a default prefix, namespace pairs for pretty output
default_ns = Namespace(
    "http://example.org/BIM_lighting_properties_standards/CEN_TC_274_Ontology#")
store.bind("", default_ns)

# create a property hasID for the human readable ID
prop_id = URIRef(default_ns + "hasID")
store.add((prop_id, RDF.type, RDF.Property))

# table 0: Table 01:  Mechanical properties
create_properties(document.tables[0], "MechanicalProperties")
# table 1: Table 02:  Electrical properties
create_properties(document.tables[1], "ElectricalProperties")
# table 2: Table 03:  Emergency lighting properties
create_properties(document.tables[2], "EmergencyLightingProperties")
# table 3: Table 04:  Photometric properties
create_properties(document.tables[3], "PhotometricProperties")
# table 4: Table 05:  Sensing device properties
create_properties(document.tables[4], "SensingDeviceProperties")
# table 5: Table 06:  Mounting & Accessory properties
create_properties(document.tables[5], "MountingAccessoryProperties")
# table 6: Table 07:  Marketing properties
create_properties(document.tables[6], "MarketingProperties")
# table 7: Table 08:  Operations & Maintenance properties
create_properties(document.tables[7], "OperationsMaintenanceProperties")


# Serialize as Turtle
rdffile = open('CEN_TC_274_Ontology.ttl', 'wb')
rdffile.write(store.serialize(format="turtle"))

rdffile.close()
